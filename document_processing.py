from pathlib import Path
import streamlit as st
from langchain.chains.conversational_retrieval.base import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders.pdf import PyPDFLoader
from langchain_community.vectorstores.faiss import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai.embeddings import OpenAIEmbeddings
from langchain_openai.chat_models import ChatOpenAI
from dotenv import load_dotenv, find_dotenv
from configs import *
from langchain_community.document_loaders import UnstructuredHTMLLoader
from langchain.docstore.document import Document
from langchain_core.documents import Document

# Importações adicionais
from io import BytesIO
import docx
import fitz  # PyMuPDF
from PIL import Image
import io
import requests
import pdfplumber
import os
from ebooklib import epub
import ebooklib
from bs4 import BeautifulSoup
import pandas as pd
import tempfile
import openai

_ = load_dotenv(find_dotenv())

PASTA_ARQUIVOS = Path(__file__).parent / 'arquivos'
# Declaração das variáveis globais
documentos = []
images = []
tables = []

def descrever_imagem_com_gpt4(image):
    # Enviando a imagem para o GPT-4 com suporte a visão
    # Isso requer que você tenha uma função ou uma API que aceite imagens e retorne descrições
    response = openai.Image.create(file=image, model="gpt-4-vision")
    descricao_imagem = response['data']['description']  # A descrição retornada pelo modelo
    return descricao_imagem

def transcrever_imagens_com_gpt4(images_pdf):
    transcricoes = []
    for i, image in enumerate(images_pdf):
        descricao_imagem = descrever_imagem_com_gpt4(image)
        transcricoes.append(f"Descrição da imagem {i + 1}: {descricao_imagem}")
    return "\n\n".join(transcricoes)

def tabelas_para_string(tables):
    tabelas_texto = ""
    for i, table in enumerate(tables):
        if table:
            for row in table:
            # Converte cada DataFrame para string e adiciona ao texto final
                tabelas_texto += f"\nTabela {i + 1}:\n{row}\n"
        else:
            tabelas_texto += f"\nTabela {i + 1}: O objeto não é um DataFrame\n"
    return tabelas_texto

def importacao_documentos():
    global documentos, images, tables  # Declarar as variáveis como globais para uso fora da função

    documentos = []  # Reinicia as variáveis a cada execução da função para evitar dados duplicados
    images = []
    tables = []
    
    for arquivo in PASTA_ARQUIVOS.iterdir():
        if arquivo.suffix.lower() == '.pdf':

            texto, images_pdf, tables_pdf = extract_content_from_pdf(arquivo)
            #st.write(f"Conteúdo do PDF {arquivo.name}:")
            # transcricoes_imagens = transcrever_imagens_com_gpt4(images_pdf)
            #st.write(texto)
            images.extend(images_pdf)
            tables.extend(tables_pdf)

            tabelas_texto = tabelas_para_string(tables_pdf)

            documentos.append(Document(page_content=f"{texto}", metadata={"source": arquivo.name}))

        elif arquivo.suffix.lower() == '.docx':
            texto = extract_text_from_docx(arquivo)
            documentos.append(Document(page_content=f"{texto}", metadata={"source": arquivo.name}))
            
        elif arquivo.suffix.lower() == '.epub':
            texto = extract_text_from_epub(arquivo)
            documentos.append(Document(page_content=f"{texto}", metadata={"source": arquivo.name}))
            
        elif arquivo.suffix.lower() in ['.xlsx', '.xls']:
            texto = extract_text_from_excel(arquivo)
            documentos.append(Document(page_content=f"{texto}", metadata={"source": arquivo.name}))
            
        elif arquivo.suffix.lower() == '.html':
            #st.write(arquivo)
            texto = load_html_content(arquivo)
            documentos.append(Document(page_content=f"{texto}", metadata={"source": arquivo.name}))
            
    return documentos, images, tables

# A partir de agora, você pode acessar as variáveis `documentos`, `images` e `tables` fora da função

def extract_content_from_pdf(pdf_path_or_file):
    # Check if the input is a URL or a file
    if isinstance(pdf_path_or_file, (str, Path)):
        # If it's a path or URL, handle it
        if str(pdf_path_or_file).startswith("http://") or str(pdf_path_or_file).startswith("https://"):
            # If the input is a URL, download the PDF
            response = requests.get(pdf_path_or_file)
            response.raise_for_status()  # Check for errors
            pdf_stream = io.BytesIO(response.content)
        else:
            # If the input is a file path, open it
            with open(pdf_path_or_file, "rb") as f:
                pdf_stream = io.BytesIO(f.read())
    else:
        # If it's a file-like object, assume it's already open
        pdf_stream = io.BytesIO(pdf_path_or_file.read())

    doc = fitz.open(stream=pdf_stream)
    images = []
    texts = []
    tables = []

    with pdfplumber.open(pdf_stream) as pdf:
        for page_num, page in enumerate(doc, start=1):
            
            text = page.get_text()
            if text:  # Verifica se o texto foi extraído
                page_info = f"\n--- Página {page_num} ---\n"
                texts.append(page_info + text)
            # Extract text from the page
           
            texts.append(text)

            # Extract images from the page
            for img in page.get_images(full=True):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image = Image.open(io.BytesIO(image_bytes))
                images.append(image)

            # Extract tables from the page using pdfplumber
            plumber_page = pdf.pages[page_num - 1]
            page_tables = plumber_page.extract_tables()
            if page_tables:
                tables.append(page_tables)

    return texts, images, tables

def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        if not text:
            raise ValueError("O arquivo DOCX está vazio ou não contém texto extraível.")
        return text
    except Exception as e:
        raise RuntimeError(f"Erro ao tentar extrair o texto do arquivo DOCX: {e}")

def extract_text_from_epub(file):
    try:
        book = epub.read_epub(file)
        text = ""
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            html_content = item.get_body_content().decode()
            text += convert_html_to_text(html_content)
        if not text.strip():
            raise ValueError("O arquivo EPUB está vazio ou não contém texto extraível.")
        return text
    except Exception as e:
        raise RuntimeError(f"Erro ao tentar extrair o texto do arquivo EPUB: {e}")

def convert_html_to_text(html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    return soup.get_text(separator="\n")

def extract_text_from_excel(file):
    try:
        df = pd.read_excel(file)
        if df.empty:
            raise ValueError("O arquivo Excel está vazio ou não contém dados.")
        return df.to_string()
    except Exception as e:
        raise RuntimeError(f"Erro ao tentar extrair o texto do arquivo Excel: {e}")

#----------------------------

def load_html_content(html_path_or_url): 

    if html_path_or_url:
        # Carregar o arquivo HTML
        with open(html_path_or_url, "r", encoding="utf-8") as file:
            conteudo_html = file.read()
        # Analisar o conteúdo HTML com BeautifulSoup
        soup = BeautifulSoup(conteudo_html, "html.parser")

        # Exemplo: extrair texto da página
        texto = soup.get_text()
    else:
        print("Erro")

    return texto

# ------------------------------

def split_de_documentos(documentos):

    recur_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2500,
        chunk_overlap=250,
        separators=["/n\n", "\n", ".", " ", ""]
    )
    documentos = recur_splitter.split_documents(documentos)
    for i, doc in enumerate(documentos):
        doc.metadata['source'] = doc.metadata['source'].split('/')[-1]
        doc.metadata['doc_id'] = i
    return documentos, doc.metadata['source']

def cria_vector_store(documentos, doc_names):
    embedding_model = OpenAIEmbeddings()
    vector_store = FAISS.from_documents(
        documents=documentos,
        embedding=embedding_model
    )

    caminho_arquivo = f"faiss_index_{doc_names}"
    vector_store.save_local(caminho_arquivo)
    return vector_store


def cria_chain_conversa():

    documentos, a, b = importacao_documentos()

    documentos, d = split_de_documentos(documentos)

    vector_store = cria_vector_store(documentos, d)

    chat = ChatOpenAI(model=get_config('model_name'))
    memory = ConversationBufferMemory(
        return_messages=True,
        memory_key='chat_history',
        output_key='answer'
    )
    retriever = vector_store.as_retriever(
        search_type=get_config('retrieval_search_type'),
        search_kwargs=get_config('retrieval_kwargs')
    )
    prompt = PromptTemplate.from_template(get_config('prompt'))
    chat_chain = ConversationalRetrievalChain.from_llm(
        llm=chat,
        memory=memory,
        retriever=retriever,
        return_source_documents=True,
        verbose=True,
        combine_docs_chain_kwargs={'prompt': prompt}
    )

    st.session_state['chain'] = chat_chain

    return documentos, a, b, d